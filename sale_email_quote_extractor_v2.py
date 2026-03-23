# -*- coding: utf-8 -*-
"""
销售邮箱报价信息提取（增强版）- 支持附件提取
登录 sale@aeroedgeglobal.com，提取 2026-03-01 以来的航材供应商报价邮件
支持提取 Excel、PDF、Word、CSV 附件中的报价信息
"""

import imaplib
import email
from email.header import decode_header
import sys
import re
from datetime import datetime
import csv
import os
import tempfile
from config import (
    IMAP_CONFIG, 
    DEFAULT_OUTPUT_DIR, 
    SUPPLIER_QUOTE_FILE,
    OUTPUT_ENCODING
)

# 第三方库导入
try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

sys.stdout.reconfigure(encoding='utf-8')

# IMAP 配置（从 config 模块导入）
IMAP_SERVER = IMAP_CONFIG["server"]
IMAP_PORT = IMAP_CONFIG["port"]
USERNAME = IMAP_CONFIG["username"]
# 注意：密码应从安全存储获取，此处为示例
AUTH_CODE = "A4D8%b3x6FHAH45d"

# 输出文件路径（使用配置的默认路径）
OUTPUT_FILE = SUPPLIER_QUOTE_FILE

# 临时目录用于保存附件
TEMP_DIR = tempfile.mkdtemp(prefix="email_attachments_")


def decode_mime_words(s):
    """解码 MIME 编码的字符串"""
    if not s:
        return ""
    decoded = decode_header(s)
    return ''.join(
        part.decode(encoding or 'utf-8', errors='ignore') if isinstance(part, bytes) else part
        for part, encoding in decoded
    )


def parse_price(text):
    """从文本中提取价格信息"""
    if not text:
        return None
    # 价格模式：$XXX, XXX.XX 或 USD XXX 或 EUR XXX
    patterns = [
        r'\$[\d,]+\.?\d*',
        r'USD\s*[\d,]+\.?\d*',
        r'EUR\s*[\d,]+\.?\d*',
        r'CNY\s*[\d,]+\.?\d*',
        r'price[:\s]*\$?[\d,]+\.?\d*',
        r'unit\s*price[:\s]*\$?[\d,]+\.?\d*',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(0)
    return None


def parse_condition(text):
    """从文本中提取条件"""
    if not text:
        return ""
    conditions = ['SV', 'NS', 'NE', 'OH', 'AR', 'FN', 'AS', 'DIST']
    text_upper = text.upper()
    for cond in conditions:
        if re.search(rf'\b{cond}\b', text_upper):
            return cond
    return ""


def parse_quantity(text):
    """从文本中提取数量"""
    if not text:
        return ""
    patterns = [
        r'Qty[:\s]*(\d+)',
        r'Quantity[:\s]*(\d+)',
        r'(\d+)\s*pcs?',
        r'(\d+)\s*units?',
        r'(\d+)\s*EA',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1)
    return ""


def extract_part_number(text):
    """从文本中提取零件号"""
    if not text:
        return ""
    patterns = [
        r'PN[:\s]*([A-Z0-9\-]+)',
        r'Part\s*(?:Number)?[:\s]*([A-Z0-9\-]+)',
        r'P/N[:\s]*([A-Z0-9\-]+)',
        r'#([A-Z0-9\-]+)',
        r'([A-Z]{2,}\d{3,}[-\d]*)',
        r'(\d{2,}-?\d{4,}-?\d{0,})',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            pn = match.group(1)
            # 过滤掉太短的匹配
            if len(pn) >= 5:
                return pn
    return ""


def extract_quotes_from_text(text, source_info):
    """从文本中提取报价信息"""
    quotes = []
    lines = text.split('\n')
    
    current_pn = None
    current_desc = None
    current_price = None
    current_condition = None
    current_qty = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 尝试提取零件号
        pn = extract_part_number(line)
        if pn:
            # 保存之前的记录
            if current_pn:
                quotes.append({
                    'Date': source_info.get('Date', ''),
                    'From': source_info.get('From', ''),
                    'Subject': source_info.get('Subject', ''),
                    'PartNumber': current_pn,
                    'Description': current_desc or '',
                    'Price': current_price or '',
                    'Condition': current_condition or '',
                    'Qty': current_qty or '',
                    'Source': source_info.get('Source', 'Email Body'),
                    'Notes': ''
                })
            
            current_pn = pn
            current_desc = line.replace(pn, '').strip()
            current_price = parse_price(line)
            current_condition = parse_condition(line)
            current_qty = parse_quantity(line)
        else:
            # 尝试从当前行提取价格、条件、数量
            if current_pn:
                if not current_price:
                    current_price = parse_price(line)
                if not current_condition:
                    current_condition = parse_condition(line)
                if not current_qty:
                    current_qty = parse_quantity(line)
    
    # 保存最后一条记录
    if current_pn:
        quotes.append({
            'Date': source_info.get('Date', ''),
            'From': source_info.get('From', ''),
            'Subject': source_info.get('Subject', ''),
            'PartNumber': current_pn,
            'Description': current_desc or '',
            'Price': current_price or '',
            'Condition': current_condition or '',
            'Qty': current_qty or '',
            'Source': source_info.get('Source', 'Email Body'),
            'Notes': ''
        })
    
    return quotes


def extract_from_excel(file_path, source_info):
    """从 Excel 文件中提取报价信息"""
    quotes = []
    
    if not HAS_OPENPYXL and not HAS_PANDAS:
        print(f"  [WARN] 缺少 openpyxl 或 pandas 库，跳过 Excel 文件")
        return quotes
    
    try:
        if HAS_PANDAS:
            # 使用 pandas 读取所有 sheet
            xls = pd.ExcelFile(file_path)
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                # 转换为文本
                text = df.to_string()
                extracted = extract_quotes_from_text(text, source_info)
                quotes.extend(extracted)
        elif HAS_OPENPYXL:
            # 使用 openpyxl 读取
            wb = openpyxl.load_workbook(file_path, data_only=True)
            for ws in wb.worksheets:
                text = ""
                for row in ws.iter_rows(values_only=True):
                    row_text = " ".join([str(cell) if cell else "" for cell in row])
                    text += row_text + "\n"
                extracted = extract_quotes_from_text(text, source_info)
                quotes.extend(extracted)
        
        print(f"  [OK] 从 Excel 提取到 {len(quotes)} 条报价")
    except Exception as e:
        print(f"  [ERROR] Excel 提取失败：{e}")
    
    return quotes


def extract_from_pdf(file_path, source_info):
    """从 PDF 文件中提取报价信息"""
    quotes = []
    
    if not HAS_PDFPLUMBER:
        print(f"  [WARN] 缺少 pdfplumber 库，跳过 PDF 文件")
        return quotes
    
    try:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        extracted = extract_quotes_from_text(text, source_info)
        quotes.extend(extracted)
        print(f"  [OK] 从 PDF 提取到 {len(quotes)} 条报价")
    except Exception as e:
        print(f"  [ERROR] PDF 提取失败：{e}")
    
    return quotes


def extract_from_word(file_path, source_info):
    """从 Word 文件中提取报价信息"""
    quotes = []
    
    if not HAS_DOCX:
        print(f"  [WARN] 缺少 python-docx 库，跳过 Word 文件")
        return quotes
    
    try:
        doc = docx.Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # 也提取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join([cell.text for cell in row.cells])
                text += row_text + "\n"
        
        extracted = extract_quotes_from_text(text, source_info)
        quotes.extend(extracted)
        print(f"  [OK] 从 Word 提取到 {len(quotes)} 条报价")
    except Exception as e:
        print(f"  [ERROR] Word 提取失败：{e}")
    
    return quotes


def extract_from_csv(file_path, source_info):
    """从 CSV 文件中提取报价信息"""
    quotes = []
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        extracted = extract_quotes_from_text(text, source_info)
        quotes.extend(extracted)
        print(f"  [OK] 从 CSV 提取到 {len(quotes)} 条报价")
    except Exception as e:
        print(f"  [ERROR] CSV 提取失败：{e}")
    
    return quotes


def save_attachment(part, temp_dir):
    """保存附件到临时目录"""
    filename = part.get_filename()
    if not filename:
        return None
    
    filename = decode_mime_words(filename)
    
    # 清理文件名
    filename = re.sub(r'[^\w\.\-]', '_', filename)
    
    filepath = os.path.join(temp_dir, filename)
    
    payload = part.get_payload(decode=True)
    if payload:
        with open(filepath, 'wb') as f:
            f.write(payload)
        return filepath
    return None


def process_attachments(msg, source_info, temp_dir):
    """处理邮件的所有附件"""
    all_quotes = []
    
    if not msg.is_multipart():
        return all_quotes
    
    for part in msg.walk():
        content_type = part.get_content_type()
        content_disposition = str(part.get("Content-Disposition", ""))
        
        # 检查是否是附件
        if 'attachment' in content_disposition.lower():
            filename = part.get_filename()
            if filename:
                filename = decode_mime_words(filename).lower()
                print(f"\n  发现附件：{filename}")
                
                # 保存附件
                filepath = save_attachment(part, temp_dir)
                if filepath:
                    print(f"  已保存：{filepath}")
                    
                    # 根据文件类型提取
                    if filename.endswith(('.xlsx', '.xls')):
                        quotes = extract_from_excel(filepath, source_info)
                        all_quotes.extend(quotes)
                    elif filename.endswith('.pdf'):
                        quotes = extract_from_pdf(filepath, source_info)
                        all_quotes.extend(quotes)
                    elif filename.endswith(('.docx', '.doc')):
                        quotes = extract_from_word(filepath, source_info)
                        all_quotes.extend(quotes)
                    elif filename.endswith('.csv'):
                        quotes = extract_from_csv(filepath, source_info)
                        all_quotes.extend(quotes)
                    else:
                        print(f"  [SKIP] 不支持的附件格式：{filename}")
    
    return all_quotes


print("="*80)
print("=== 销售邮箱报价信息提取（增强版 - 支持附件）===")
print(f"开始时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
print(f"输出目录：{DEFAULT_OUTPUT_DIR}")
print(f"临时目录：{TEMP_DIR}")
print("="*80)

# 检查依赖库
print("\n依赖库检查:")
print(f"  openpyxl: {'✅ 已安装' if HAS_OPENPYXL else '❌ 未安装'}")
print(f"  pdfplumber: {'✅ 已安装' if HAS_PDFPLUMBER else '❌ 未安装'}")
print(f"  python-docx: {'✅ 已安装' if HAS_DOCX else '❌ 未安装'}")
print(f"  pandas: {'✅ 已安装' if HAS_PANDAS else '❌ 未安装'}")
print("="*80)

# 连接 IMAP 服务器
print("\n[1/6] 连接 IMAP 服务器...")
try:
    mail = imaplib.IMAP4_SSL(IMAP_SERVER, IMAP_PORT)
    mail.login(USERNAME, AUTH_CODE)
    print("[OK] 登录成功")
except Exception as e:
    print(f"[ERROR] 登录失败：{e}")
    sys.exit(1)

# 选择收件箱
print("\n[2/6] 选择收件箱...")
mail.select("INBOX")

# 搜索 2026-03-01 以来的邮件
print("\n[3/6] 搜索 2026-03-01 以来的邮件...")
status, messages = mail.search(None, '(SINCE "01-Mar-2026")')

if status != "OK":
    print("[ERROR] 搜索失败")
    sys.exit(1)

email_ids = messages[0].split()
print(f"[OK] 找到 {len(email_ids)} 封邮件")

# 提取报价信息
print("\n[4/6] 提取报价信息（含附件）...")
all_quotes = []
email_count = 0
attachment_count = 0

# 关键词过滤
keywords = ['quote', 'quotation', 'price', 'offer', 'rfq', 'pn', 'part number']

for i, eid in enumerate(email_ids, 1):
    status, msg_data = mail.fetch(eid, "(RFC822)")
    if status != "OK":
        continue
    
    for response_part in msg_data:
        if isinstance(response_part, tuple):
            msg = email.message_from_bytes(response_part[1])
            
            subject = decode_mime_words(msg.get("Subject", ""))
            from_ = decode_mime_words(msg.get("From", ""))
            date_str = msg.get("Date", "")
            
            try:
                date_obj = email.utils.parsedate_to_datetime(date_str)
                formatted_date = date_obj.strftime("%Y-%m-%d %H:%M")
            except:
                formatted_date = date_str
            
            # 检查是否包含报价相关关键词
            subject_lower = subject.lower()
            if not any(kw in subject_lower for kw in keywords):
                continue
            
            # 准备源信息
            source_info = {
                'Date': formatted_date,
                'From': from_[:50],
                'Subject': subject[:50],
                'Source': 'Email Body'
            }
            
            # 提取正文
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    ctype = part.get_content_type()
                    cdispo = str(part.get("Content-Disposition"))
                    if ctype == "text/plain" and "attachment" not in cdispo:
                        try:
                            body = part.get_payload(decode=True).decode('utf-8', errors='ignore')
                        except:
                            pass
                        break
                else:
                    try:
                        body = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
                    except:
                        pass
            
            # 从正文提取
            if body:
                quotes = extract_quotes_from_text(body, source_info)
                all_quotes.extend(quotes)
                if quotes:
                    email_count += 1
                    print(f"\n邮件 {email_count}: {subject[:60]}")
                    print(f"  发件人：{from_[:50]}")
                    print(f"  日期：{formatted_date}")
                    print(f"  提取到 {len(quotes)} 条报价（正文）")
            
            # 处理附件
            attachment_quotes = process_attachments(msg, source_info, TEMP_DIR)
            if attachment_quotes:
                all_quotes.extend(attachment_quotes)
                attachment_count += 1
                print(f"  提取到 {len(attachment_quotes)} 条报价（附件）")

print(f"\n\n总计：{email_count} 封报价邮件，{attachment_count} 封含附件，{len(all_quotes)} 条报价记录")

# 导出 CSV
print("\n[5/6] 导出到 CSV...")

if all_quotes:
    fieldnames = ['Date', 'From', 'Subject', 'PartNumber', 'Description', 'Price', 'Condition', 'Qty', 'Source', 'Notes']
    with open(OUTPUT_FILE, 'w', newline='', encoding=OUTPUT_ENCODING) as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(all_quotes)
    
    print(f"[OK] 已导出到：{OUTPUT_FILE}")
    
    # 统计信息
    print("\n[6/6] 统计信息...")
    print("\n=== 报价统计 ===")
    
    # 按供应商统计
    supplier_stats = {}
    for q in all_quotes:
        supplier = q['From'].split('<')[-1].split('>')[0] if '<' in q['From'] else q['From']
        supplier_stats[supplier] = supplier_stats.get(supplier, 0) + 1
    
    print("\n按供应商统计 (Top 10):")
    for supplier, count in sorted(supplier_stats.items(), key=lambda x: x[1], reverse=True)[:10]:
        print(f"  {supplier}: {count} 条")
    
    # 按条件统计
    cond_stats = {}
    for q in all_quotes:
        cond = q['Condition'] or 'N/A'
        cond_stats[cond] = cond_stats.get(cond, 0) + 1
    
    print("\n按条件统计:")
    for cond, count in sorted(cond_stats.items(), key=lambda x: x[1], reverse=True):
        print(f"  {cond}: {count} 条")
    
    # 按来源统计
    source_stats = {}
    for q in all_quotes:
        source = q['Source']
        source_stats[source] = source_stats.get(source, 0) + 1
    
    print("\n按来源统计:")
    for source, count in sorted(source_stats.items(), key=lambda x: x[1], reverse=True):
        print(f"  {source}: {count} 条")
else:
    print("[WARN] 未找到报价记录")

# 清理临时文件
print("\n清理临时文件...")
try:
    import shutil
    shutil.rmtree(TEMP_DIR)
    print(f"[OK] 已清理：{TEMP_DIR}")
except Exception as e:
    print(f"[WARN] 清理失败：{e}")

# 关闭连接
mail.close()
mail.logout()

print(f"\n结束时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
print("=== 任务完成 ===")
